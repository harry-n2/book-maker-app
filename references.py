"""参照ソース（URL / ファイル / 画像 / NotebookLM）の取り込みモジュール。

各参照を統一フォーマット {"label": ..., "kind": ..., "content": ...} で返す。
generator.py はこれを context として Gemini プロンプトに組み込む。
"""

from __future__ import annotations

import base64
import ipaddress
import mimetypes
import re
import socket
from pathlib import Path
from typing import NotRequired, TypedDict
from urllib.parse import urlparse

import requests


MAX_BYTES_PER_SOURCE = 200_000  # 1ソースあたり最大文字バイト数（誤った巨大ファイル対策）
USER_AGENT = "Mozilla/5.0 (compatible; BookMakerApp/1.0)"
REQUEST_TIMEOUT = 15

# 参照本文がこれ未満なら「短すぎ警告」を出す。動的レンダリングや権限不足で本文が取れていない可能性。
MIN_CONTENT_CHARS_FOR_OK = 500

# Google Docs / Sheets の編集 URL を export 形式に書き換える正規表現
_GDOC_RE = re.compile(r"^(https?://docs\.google\.com)/document/d/([a-zA-Z0-9_-]+)(?:/.*)?$")
_GSHEET_RE = re.compile(r"^(https?://docs\.google\.com)/spreadsheets/d/([a-zA-Z0-9_-]+)(?:/.*)?$")
_GSLIDE_RE = re.compile(r"^(https?://docs\.google\.com)/presentation/d/([a-zA-Z0-9_-]+)(?:/.*)?$")


class Reference(TypedDict):
    label: str
    kind: str
    content: str
    warning: NotRequired[str]


# ---------------------------------------------------------------------------
# URL
# ---------------------------------------------------------------------------


def _is_safe_url(url: str) -> bool:
    """SSRF 対策：内部 IP やプライベートネットワークを弾く。"""
    try:
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"}:
            return False
        host = parsed.hostname
        if not host:
            return False
        try:
            ip = ipaddress.ip_address(host)
        except ValueError:
            try:
                resolved = socket.gethostbyname(host)
                ip = ipaddress.ip_address(resolved)
            except (socket.gaierror, ValueError):
                return False
        if ip.is_private or ip.is_loopback or ip.is_reserved or ip.is_link_local:
            return False
        return True
    except Exception:
        return False


def _rewrite_google_url(url: str) -> tuple[str, str | None]:
    """Google Docs / Sheets / Slides の編集 URL を export 形式に書き換える。

    Slides は txt エクスポート不可なので警告を返すために種別だけ識別する。

    戻り値: (書き換え後 URL, 種別: 'docs' | 'sheets' | 'slides' | None)
    """
    m = _GDOC_RE.match(url)
    if m:
        return f"{m.group(1)}/document/d/{m.group(2)}/export?format=txt", "docs"
    m = _GSHEET_RE.match(url)
    if m:
        return f"{m.group(1)}/spreadsheets/d/{m.group(2)}/export?format=csv", "sheets"
    m = _GSLIDE_RE.match(url)
    if m:
        # Slides はテキストエクスポート不可。本文取得を試みるが警告対象になる。
        return url, "slides"
    return url, None


def _detect_url_warning(text: str, gtype: str | None) -> str | None:
    """取得本文に警告すべき特徴があれば文言を返す。"""
    if gtype == "slides":
        return (
            "Google Slides は直接テキスト抽出ができません。"
            "スライドを PDF または txt にダウンロードして「ファイル」欄に添付してください"
        )
    if not text or not text.strip():
        return "本文がほぼ空です。URL の公開設定または対応形式を見直してください"
    stripped_len = len(text.strip())
    if stripped_len < MIN_CONTENT_CHARS_FOR_OK:
        return (
            f"取得本文が {stripped_len} 文字と短いため、本文が取れていない可能性があります。"
            "ダウンロード版を「ファイル」欄に添付するか、テキスト直貼り欄に本文をペーストしてください"
        )
    head = text[:600]
    if "JavaScript" in head and ("有効" in head or "enable" in head.lower()):
        return (
            "JavaScript レンダリングが必要なページのため本文が取れていません。"
            "Google Docs であれば自動で /export?format=txt を試みますが失敗した場合は、"
            "ダウンロードしたファイルを添付するかテキスト直貼り欄を使ってください"
        )
    if ("ログイン" in head or "Sign in" in head) and ("アカウント" in head or "account" in head.lower()):
        return (
            "ログインが必要なページの可能性があります。"
            "公開設定にするか、コンテンツをエクスポートして添付してください"
        )
    return None


def fetch_url(url: str) -> Reference:
    url = url.strip()
    if not _is_safe_url(url):
        return {
            "label": url,
            "kind": "url",
            "content": "[取得失敗：URL が無効または内部ネットワーク]",
            "warning": "URL が無効です（http/https のみ・公開ホストのみ対応）",
        }

    # Google Docs / Sheets / Slides は export URL に書き換えてから取得
    fetch_target, gtype = _rewrite_google_url(url)
    try:
        response = requests.get(
            fetch_target,
            headers={"User-Agent": USER_AGENT},
            timeout=REQUEST_TIMEOUT,
            allow_redirects=True,
        )
        if response.status_code in (401, 403, 404):
            warn = (
                "Google ドキュメントの公開設定の問題で取得失敗（HTTP "
                f"{response.status_code}）。「リンクを知っている全員に閲覧可」に変更するか、"
                "ダウンロードしたファイルを添付してください"
                if gtype
                else f"取得失敗（HTTP {response.status_code}）。ページの公開設定または URL を確認してください"
            )
            return {
                "label": url,
                "kind": "url",
                "content": f"[取得失敗：HTTP {response.status_code}]",
                "warning": warn,
            }
        response.raise_for_status()

        # Google エクスポート版はそのままテキスト（docs=txt / sheets=csv）。
        # 通常 URL は trafilatura で本文抽出。
        if gtype in ("docs", "sheets"):
            text = response.text
        else:
            text = _extract_main_text(response.text, url)

        text = text[:MAX_BYTES_PER_SOURCE]
        ref: Reference = {"label": url, "kind": "url", "content": text}
        warn = _detect_url_warning(text, gtype)
        if warn:
            ref["warning"] = warn
        return ref
    except Exception as exc:  # noqa: BLE001
        return {
            "label": url,
            "kind": "url",
            "content": f"[取得失敗：{exc}]",
            "warning": f"取得時に例外が発生しました：{exc}",
        }


def _extract_main_text(html: str, url: str) -> str:
    try:
        import trafilatura

        extracted = trafilatura.extract(
            html,
            url=url,
            include_comments=False,
            include_tables=True,
            favor_recall=True,
        )
        if extracted:
            return extracted
    except Exception:
        pass
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
            tag.decompose()
        text = soup.get_text("\n", strip=True)
        return re.sub(r"\n{3,}", "\n\n", text)
    except Exception:
        return html


# ---------------------------------------------------------------------------
# ファイル（PDF / DOCX / MD / TXT / CSV）
# ---------------------------------------------------------------------------


def extract_file(path: Path, original_name: str | None = None) -> Reference:
    name = original_name or path.name
    ext = Path(name).suffix.lower()
    label = f"file: {name}"
    try:
        if ext == ".pdf":
            content = _extract_pdf(path)
        elif ext == ".docx":
            content = _extract_docx(path)
        elif ext in {".md", ".markdown", ".txt", ".csv", ".json", ".yml", ".yaml"}:
            content = path.read_text(encoding="utf-8", errors="replace")
        else:
            content = f"[未対応の拡張子：{ext}]"
        return {"label": label, "kind": "file", "content": content[:MAX_BYTES_PER_SOURCE]}
    except Exception as exc:  # noqa: BLE001
        return {"label": label, "kind": "file", "content": f"[読み込み失敗：{exc}]"}


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
            parts.append(f"--- Page {i+1} ---\n{text.strip()}")
        except Exception as exc:  # noqa: BLE001
            parts.append(f"--- Page {i+1} ---\n[抽出失敗：{exc}]")
        if sum(len(p) for p in parts) > MAX_BYTES_PER_SOURCE:
            break
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 画像（Gemini Vision で説明テキスト生成）
# ---------------------------------------------------------------------------


def analyze_image(path: Path, original_name: str, api_key: str, model: str = "") -> Reference:
    label = f"image: {original_name}"
    try:
        import google.generativeai as genai
        # 画像読取は現行Stableのマルチモーダルモデルを使う（旧 gemini-2.0-flash は Shut down）。
        # generator.resolve_model で停止モデルの自己修復＋env上書きを一元化。
        from generator import resolve_model

        genai.configure(api_key=api_key)
        m = genai.GenerativeModel(resolve_model(model))
        mime, _ = mimetypes.guess_type(original_name)
        if not mime:
            mime = "image/png"
        with open(path, "rb") as fp:
            data = fp.read()
        prompt = (
            "この画像の内容を、書籍の素材として活用できる形で詳しく日本語で説明してください。"
            "図解の場合は構造、テキストの場合は内容、写真の場合は被写体と文脈を抽出してください。"
        )
        response = m.generate_content([
            prompt,
            {"mime_type": mime, "data": data},
        ])
        return {
            "label": label,
            "kind": "image",
            "content": (response.text or "").strip()[:MAX_BYTES_PER_SOURCE],
        }
    except Exception as exc:  # noqa: BLE001
        return {"label": label, "kind": "image", "content": f"[画像解析失敗：{exc}]"}


# ---------------------------------------------------------------------------
# NotebookLM 連携（公開ノートのコンテンツを URL 取得で吸い込む）
# ---------------------------------------------------------------------------


def fetch_notebooklm(url: str) -> Reference:
    """NotebookLM の共有 URL から取得を試みる。

    NotebookLM には公式 API が公開されていないため、共有 URL からの取得は
    Google 側のレンダリング仕様に依存する。動的ロード分は取得できない場合がある。
    取得できた範囲をテキストとして返す。
    """
    url = url.strip()
    label = f"notebooklm: {url}"
    if "notebooklm.google.com" not in url:
        return {
            "label": label,
            "kind": "notebooklm",
            "content": "[NotebookLM の URL ではありません]",
        }
    if not _is_safe_url(url):
        return {"label": label, "kind": "notebooklm", "content": "[URL が無効]"}
    try:
        response = requests.get(
            url,
            headers={"User-Agent": USER_AGENT},
            timeout=REQUEST_TIMEOUT,
            allow_redirects=True,
        )
        response.raise_for_status()
        text = _extract_main_text(response.text, url)
        if not text or len(text) < 200:
            text = (
                "[NotebookLM はクライアントサイド・レンダリングのため、"
                "公開ノート本文を直接取得できない場合があります。"
                "代わりに NotebookLM 内で『エクスポート』機能で書き出した Markdown / PDF を"
                "ファイル参照としてアップロードしてください]"
            )
        return {"label": label, "kind": "notebooklm", "content": text[:MAX_BYTES_PER_SOURCE]}
    except Exception as exc:  # noqa: BLE001
        return {"label": label, "kind": "notebooklm", "content": f"[取得失敗：{exc}]"}


# ---------------------------------------------------------------------------
# 集約ヘルパ
# ---------------------------------------------------------------------------


def fetch_pasted_text(text: str, label: str = "テキスト直貼り") -> Reference:
    """ユーザーが入力フォームに直接ペーストした本文を参照ソース化する。"""
    text = (text or "").strip()
    if not text:
        return {
            "label": label,
            "kind": "pasted",
            "content": "",
            "warning": "貼り付けテキストが空です",
        }
    if len(text) > MAX_BYTES_PER_SOURCE:
        text = text[:MAX_BYTES_PER_SOURCE]
    ref: Reference = {"label": label, "kind": "pasted", "content": text}
    if len(text) < 200:
        ref["warning"] = (
            f"貼り付けテキストが {len(text)} 文字と短いため、本書の素材として不十分な可能性があります"
        )
    return ref


def render_references_block(refs: list[Reference]) -> str:
    """参照を Gemini プロンプト用にフォーマット。

    v1.2: 参照ソースが本書の **核** であることを強制する。
    voice.md v7.1 規範（ハリー口調・failure_bank・中流階級KW）よりも、参照ソースの
    主題・固有名詞・章立て・コマンド/テンプレを優先させる。
    """
    if not refs:
        return ""
    lines = ["【参照ソース（本書の核として使用すること）】"]
    has_warning = False
    for i, r in enumerate(refs, 1):
        lines.append(f"--- 参照 {i}: {r['kind']} | {r['label']} ---")
        lines.append(r["content"])
        warn = r.get("warning")
        if warn:
            has_warning = True
            lines.append(f"⚠ この参照には警告があります: {warn}")
        lines.append("")
    lines.append("【参照ソース ここまで】")
    lines.append("")
    lines.append(
        "【参照ソースの優先順位ルール（厳守）】\n"
        "- 上記参照ソースの内容を本書の核として使ってください\n"
        "- タイトル・サブタイトルは参照ソースの主題から離れてはいけない\n"
        "- 章立て（H1/H2/H3）は参照ソースで言及された手順・段階・章構成を反映する\n"
        "- 各章の核メッセージ（key_message）は参照ソースの主張から派生させる\n"
        "- 本文の説明・コピペ枠の中身は、参照ソースで言及された具体コマンド・テンプレ・コード・固有名詞を最優先で使う\n"
        "- 著者プロフィール要素（営業◯年・LP◯枚・Kindle◯冊 等）は参照ソースの主題と整合する場合のみ使用し、矛盾する場合は省略する\n"
        "- ハリー個人の経験・voice.md v7.1 規範よりも、参照ソースの内容を優先する\n"
        "- 参照ソースが空または極端に短い（全体で500文字未満）と判断した場合は、"
        "ハリー定型ではなく『参照ソースが不足しているため生成できない』旨を出力に含めて短く返してください\n"
    )
    if has_warning:
        lines.append(
            "⚠ 一部の参照ソースに警告があります。警告のあるソースは内容が薄い可能性があるため、"
            "別の警告なし参照ソースを優先してください。\n"
        )
    return "\n".join(lines)
