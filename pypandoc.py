"""Small local pypandoc-compatible shim.

Vercel cannot bundle pypandoc-binary because the embedded Pandoc payload pushes
the Python function over the 500 MB limit. This module keeps the app deployable:
if a system Pandoc exists it is used, otherwise DOCX output falls back to a
simple python-docx conversion.
"""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path


def get_pandoc_path() -> str:
    path = shutil.which("pandoc")
    if not path:
        raise OSError("pandoc executable was not found")
    return path


def convert_file(
    source_file: str,
    to: str,
    outputfile: str,
    format: str | None = None,  # noqa: A002
    extra_args: list[str] | None = None,
) -> str:
    if to != "docx":
        raise ValueError("local pypandoc shim only supports docx output")

    pandoc = shutil.which("pandoc")
    if pandoc:
        cmd = [pandoc, source_file, "-o", outputfile]
        if format:
            cmd.extend(["-f", format])
        cmd.extend(extra_args or [])
        subprocess.run(cmd, check=True)
        return outputfile

    from docx import Document

    doc = Document()
    for raw in Path(source_file).read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("```"):
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(outputfile)
    return outputfile
